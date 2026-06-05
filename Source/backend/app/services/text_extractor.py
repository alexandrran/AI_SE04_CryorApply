from io import BytesIO

from fastapi import UploadFile
from pypdf import PdfReader
from docx import Document


async def extract_text_from_upload(file: UploadFile) -> str:
    content = await file.read()

    if not content:
        return ""

    filename = file.filename.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(content)

    if filename.endswith(".docx"):
        return extract_text_from_docx(content)

    return ""


def extract_text_from_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    page_text = [page.extract_text() or "" for page in reader.pages]
    return normalize_text("\n".join(page_text))


def extract_text_from_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    table_cells = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                table_cells.append(cell.text)

    return normalize_text("\n".join(paragraphs + table_cells))


def normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)
